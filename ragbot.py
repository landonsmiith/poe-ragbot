from fastapi import FastAPI
from pydantic import BaseModel
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import SystemMessage
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

import os

from dotenv import load_dotenv

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

app = FastAPI()

class QueryRequest(BaseModel):
    query: str

# Load FAISS index
embeddings = OpenAIEmbeddings(api_key=os.environ.get("OPENAI_API_KEY"))
# Rebuild FAISS from docx every time

from docx import Document

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
    return "\n".join(text)

def split_by_custom_headings(text):
    headings = {
        "Logistics": ["Safety", "Parking", "Building", "Scheduling & Meeting with Faculty and Staff"],
        "Safety": ["Fire Drill", "Fire Extinguisher Locations", "First Aid", "IAA Pantry", "Secure Facility"],
        "Building": ["Floorplan / Map", "Datawatch Card", "Lost and Found", "Lockers", "Supplies & Building Maintenance", "Kitchens"],
        "Kitchens": ["Kitchen Duty", "Food"],
        "Scheduling & Meeting with Faculty and Staff": ["Calendars", "Conference Rooms", "Email", "Contact Faculty/Staff", "Contact IT Staff", "Meet with Faculty/Staff: ScheduleOnce", "Slack", "Zoom for Meetings"],
        "Academics": ["Curriculum", "Academic Standing", "Academic Integrity & Code of Student Conduct", "Attendance", "Certifications", "Moodle"],
        "Attendance": ["Webcast Procedures for Watching Classes Live When Absent"],
        "Certifications": ["Python", "Tableau Certification", "AWS"],
        "Attire": ["Business Casual", "Business Formal", "Casual Fridays"],
        "Practicum": ["Details", "General Travel Overview", "Professional Development", "Peer Feedback"],
        "Peer Feedback": ["Self-Management", "Relationship Management", "Communication"],
        "Career Services": ["The actual job search and application process", "Corporate Relations", "Career Education", "CC: Career Conversations"],
        "Counseling": [],
        "Class of 2025 Memes": []
    }

    all_headings = set()
    for parent, children in headings.items():
        all_headings.add(parent)
        all_headings.update(children)

    sections = []
    current_heading = None
    current_content = []

    for line in text.splitlines():
        line = line.strip()

        if line in all_headings:
            if current_heading is not None and current_content:
                sections.append((current_heading, "\n".join(current_content)))
            current_heading = line
            current_content = []
        else:
            current_content.append(line)

    if current_heading is not None and current_content:
        sections.append((current_heading, "\n".join(current_content)))

    return sections


def create_faiss_index(sections):
    texts = []
    metadata = []

    for heading, content in sections:
        combined_text = f"{heading}\n{content}"
        texts.append(combined_text)
        metadata.append({"heading": heading})

    vectorstore = FAISS.from_texts(texts, embeddings, metadatas=metadata)
    return vectorstore


handbook_text = extract_text_from_docx("MSA 2025 Handbook.docx")
sections = split_by_custom_headings(handbook_text)
vectorstore = create_faiss_index(sections)
retriever = vectorstore.as_retriever()
retriever = vectorstore.as_retriever()

# System prompt to enforce handbook-only behavior
system_prompt = (
    "IMPORTANT: YOU HAVE NO ACCESS TO ANY RESOURCE EXCEPT THE MSA 2025 HANDBOOK. "
    "YOU MUST ONLY ANSWER BASED ON THE CONTENT RETRIEVED FROM THE HANDBOOK. "
    "IF THE QUESTION CANNOT BE ANSWERED FROM THE HANDBOOK, SAY: 'THE INFORMATION IS NOT AVAILABLE IN THE MSA 2025 HANDBOOK.' "
    "CITE THE RELEVANT SECTION(S) OF THE HANDBOOK IN YOUR FINAL ANSWER."
)

# ChatPromptTemplate requires 'context' because it will receive retrieved documents
prompt = ChatPromptTemplate.from_messages(
    [
        SystemMessage(content=system_prompt),
        ("user", "{context}\n\n{input}")
    ]
)

# ChatOpenAI for GPT-4 chat model
llm = ChatOpenAI(model="gpt-4", api_key=os.environ.get("OPENAI_API_KEY"))

# Create document combination and retrieval chain
stuff_chain = create_stuff_documents_chain(llm, prompt)
qa_chain = create_retrieval_chain(retriever, stuff_chain)

@app.post("/ask")
async def ask(query_request: QueryRequest):
    result = qa_chain.invoke({"input": query_request.query})
    return {"answer": result["answer"]}

# Handbook Processing Functions (Only Run When Updating Handbook)

from docx import Document
import re

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
    return "\n".join(text)

def split_by_custom_headings(text):
    headings = {
        "Logistics": [
            "Safety",
            "Parking",
            "Building",
            "Scheduling & Meeting with Faculty and Staff"
        ],
        "Safety": [
            "Fire Drill",
            "Fire Extinguisher Locations",
            "First Aid",
            "IAA Pantry",
            "Secure Facility"
        ],
        "Building": [
            "Floorplan / Map",
            "Datawatch Card",
            "Lost and Found",
            "Lockers",
            "Supplies & Building Maintenance",
            "Kitchens"
        ],
        "Kitchens": [
            "Kitchen Duty",
            "Food"
        ],
        "Scheduling & Meeting with Faculty and Staff": [
            "Calendars",
            "Conference Rooms",
            "Email",
            "Contact Faculty/Staff",
            "Contact IT Staff",
            "Meet with Faculty/Staff: ScheduleOnce",
            "Slack",
            "Zoom for Meetings"
        ],
        "Academics": [
            "Curriculum",
            "Academic Standing",
            "Academic Integrity & Code of Student Conduct",
            "Attendance",
            "Certifications",
            "Moodle"
        ],
        "Attendance": [
            "Webcast Procedures for Watching Classes Live When Absent"
        ],
        "Certifications": [
            "Python",
            "Tableau Certification",
            "AWS"
        ],
        "Attire": [
            "Business Casual",
            "Business Formal",
            "Casual Fridays"
        ],
        "Practicum": [
            "Details",
            "General Travel Overview",
            "Professional Development",
            "Peer Feedback"
        ],
        "Peer Feedback": [
            "Self-Management",
            "Relationship Management",
            "Communication"
        ],
        "Career Services": [
            "The actual job search and application process",
            "Corporate Relations",
            "Career Education",
            "CC: Career Conversations"
        ],
        "Counseling": [],
        "Class of 2025 Memes": []
    }

    all_headings = set()
    for parent, children in headings.items():
        all_headings.add(parent)
        all_headings.update(children)

    sections = []
    current_heading = None
    current_content = []

    for line in text.splitlines():
        line = line.strip()

        if line in all_headings:
            if current_heading is not None and current_content:
                sections.append((current_heading, "\n".join(current_content)))
            current_heading = line
            current_content = []
        else:
            current_content.append(line)

    if current_heading is not None and current_content:
        sections.append((current_heading, "\n".join(current_content)))

    return sections

def create_faiss_index(sections):
    texts = []
    metadata = []

    for heading, content in sections:
        combined_text = f"{heading}\n{content}"
        texts.append(combined_text)
        metadata.append({"heading": heading})

    embeddings = OpenAIEmbeddings(api_key=os.environ.get("OPENAI_API_KEY"))
    vectorstore = FAISS.from_texts(texts, embeddings, metadatas=metadata)
    vectorstore.save_local("faiss_index")

    print("FAISS index created and saved.")


# Manual Trigger to Process Handbook (Uncomment and Run Manually If Handbook Changes)
# file_path = "MSA 2025 Handbook.docx"
# handbook_text = extract_text_from_docx(file_path)
# sections = split_by_custom_headings(handbook_text)
# create_faiss_index(sections)